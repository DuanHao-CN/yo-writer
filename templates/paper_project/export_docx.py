#!/usr/bin/env python
"""将所有章节合并导出为 .docx 格式的学术论文"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
SECTIONS_DIR = os.path.join(BASE, "sections")
OUTPUT = os.path.join(BASE, "数字经济发展与企业创新质量.docx")

# 按论文顺序排列的章节文件
SECTION_FILES = [
    "abstract.md",
    "introduction.md",
    "literature_review.md",
    "related_work.md",      # 理论分析与研究假设
    "method.md",            # 研究设计
    "empirical_results.md", # 实证结果
    "discussion.md",        # 机制检验
    "experiments.md",       # 异质性分析
    "conclusion.md",        # 结论与政策启示
]

def set_doc_defaults(doc):
    """设置文档默认样式"""
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    # 设置中文字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 行距
    pf = style.paragraph_format
    pf.line_spacing = Pt(22)
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

def add_title(doc):
    """添加论文标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("数字经济发展与企业创新质量")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("——基于实质性创新与策略性创新的视角")
    run2.bold = True
    run2.font.size = Pt(15)
    run2.font.name = '黑体'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()  # 空行

def parse_md_line(line):
    """判断 markdown 行的类型"""
    stripped = line.strip()
    if not stripped:
        return ('empty', '', 0)
    if stripped.startswith('# '):
        return ('h1', stripped[2:].strip(), 1)
    if stripped.startswith('## '):
        return ('h2', stripped[3:].strip(), 2)
    if stripped.startswith('### '):
        return ('h3', stripped[4:].strip(), 3)
    if stripped.startswith('> '):
        return ('quote', stripped[2:].strip(), 0)
    if stripped.startswith('|'):
        return ('table', stripped, 0)
    if stripped.startswith('$$'):
        return ('math', stripped, 0)
    if stripped.startswith('<!--') or stripped.startswith('```'):
        return ('meta', stripped, 0)
    if stripped.startswith('[图') or stripped.startswith('[表'):
        return ('placeholder', stripped, 0)
    if stripped.startswith('- ') or stripped.startswith('* '):
        return ('list', stripped[2:].strip(), 0)
    if stripped.startswith('注：') or stripped.startswith('注:'):
        return ('note', stripped, 0)
    return ('text', stripped, 0)

def add_formatted_text(paragraph, text):
    """处理 markdown 格式（加粗、斜体）后添加到段落"""
    # 处理加粗
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:  # 奇数索引是加粗内容
            run.bold = True

def process_section(doc, filepath):
    """处理单个章节文件"""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_rows = []
    in_math = False
    math_content = []
    skip_meta = False

    for line in lines:
        line = line.rstrip('\n')
        ltype, content, level = parse_md_line(line)

        # 跳过 HTML 注释和代码块标记
        if ltype == 'meta':
            if line.strip().startswith('<!--'):
                continue
            if line.strip() == '```':
                skip_meta = not skip_meta
                continue
            continue
        if skip_meta:
            continue

        # 数学公式
        if ltype == 'math':
            if in_math:
                in_math = False
                formula = ' '.join(math_content)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(formula)
                run.font.size = Pt(11)
                run.italic = True
                math_content = []
            else:
                in_math = True
                # 检查是否是单行公式 $$ ... $$
                inner = line.strip()[2:]
                if inner.endswith('$$'):
                    inner = inner[:-2].strip()
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(inner)
                    run.font.size = Pt(11)
                    run.italic = True
                    in_math = False
            continue
        if in_math:
            math_content.append(line.strip())
            continue

        # 表格处理
        if ltype == 'table':
            if not in_table:
                in_table = True
                table_rows = []
            # 跳过分隔行
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                continue
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            if cells:
                table_rows.append(cells)
            continue
        elif in_table:
            # 表格结束，渲染
            in_table = False
            if table_rows:
                max_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=max_cols)
                table.style = 'Table Grid'
                for i, row_data in enumerate(table_rows):
                    for j, cell_text in enumerate(row_data):
                        if j < max_cols:
                            cell = table.cell(i, j)
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                paragraph.style.font.size = Pt(10)
                doc.add_paragraph()  # 表后空行
            table_rows = []

        if ltype == 'empty':
            continue

        if ltype == 'h1':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(15)
            run.font.name = '黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)

        elif ltype == 'h2':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = '黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)

        elif ltype == 'h3':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = '黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        elif ltype == 'quote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            add_formatted_text(p, content)
            p.runs[0].italic = True if p.runs else None
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

        elif ltype == 'placeholder':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.italic = True

        elif ltype == 'note':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(80, 80, 80)

        elif ltype == 'list':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            add_formatted_text(p, "  " + content)

        elif ltype == 'text':
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)  # 两个汉字缩进
            add_formatted_text(p, content)

    # 如果文件末尾还有未处理的表格
    if in_table and table_rows:
        max_cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=max_cols)
        table.style = 'Table Grid'
        for i, row_data in enumerate(table_rows):
            for j, cell_text in enumerate(row_data):
                if j < max_cols:
                    table.cell(i, j).text = cell_text

def main():
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc)

    for fname in SECTION_FILES:
        fpath = os.path.join(SECTIONS_DIR, fname)
        if os.path.exists(fpath):
            process_section(doc, fpath)
            # 章节间加一个空行
            doc.add_paragraph()
        else:
            print(f"Warning: {fpath} not found, skipping.")

    doc.save(OUTPUT)
    print(f"Done! Exported to: {OUTPUT}")

    # 统计字数
    total_chars = 0
    for p in doc.paragraphs:
        total_chars += len(p.text)
    print(f"Total characters (approx): {total_chars}")

if __name__ == '__main__':
    main()
